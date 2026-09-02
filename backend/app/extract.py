"""Content extraction from documents and web links."""
import io
import re

import httpx
from bs4 import BeautifulSoup
from fastapi import HTTPException

DOCUMENT_EXTS = {"pdf", "docx", "txt", "md", "markdown", "csv", "rtf", "vtt", "srt"}


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_subtitles(text: str) -> str:
    lines = []
    for line in text.split("\n"):
        s = line.strip()
        if not s or s.isdigit():
            continue
        if "-->" in s or s.upper().startswith("WEBVTT"):
            continue
        lines.append(s)
    return " ".join(lines)


def extract_document(data: bytes, ext: str) -> str:
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue
        return _clean("\n\n".join(pages))
    if ext == "docx":
        import docx
        document = docx.Document(io.BytesIO(data))
        blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(c.text.strip() for c in row.cells))
        return _clean("\n".join(blocks))
    if ext in {"vtt", "srt"}:
        return _clean(_strip_subtitles(data.decode("utf-8", errors="replace")))
    if ext in {"txt", "md", "markdown", "csv", "rtf"}:
        return _clean(data.decode("utf-8", errors="replace"))
    raise HTTPException(status_code=415, detail=f"Format de document non pris en charge : {ext}")


async def extract_url(url: str) -> dict:
    """Fetch a public web page / transcript link and return {'title','text'}."""
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL invalide")
    try:
        async with httpx.AsyncClient(timeout=30, follow_redirects=True,
                                     headers={"User-Agent": "Mozilla/5.0 (compatible; QuortivBot/1.0)"}) as hx:
            r = await hx.get(url)
    except Exception:  # noqa: BLE001
        raise HTTPException(status_code=502, detail="Impossible d'atteindre cette URL")
    if r.status_code >= 400:
        raise HTTPException(status_code=502, detail=f"La page a répondu {r.status_code}")

    ctype = (r.headers.get("content-type") or "").lower()
    if "text/html" not in ctype:
        for ext in ("pdf", "vtt", "srt", "txt"):
            if ext in ctype or url.lower().endswith(f".{ext}"):
                return {"title": url.rsplit("/", 1)[-1] or url, "text": extract_document(r.content, ext)}
        raise HTTPException(status_code=415, detail="Ce type de lien n'est pas exploitable")

    soup = BeautifulSoup(r.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript", "form", "svg"]):
        tag.decompose()
    title = (soup.title.string.strip() if soup.title and soup.title.string else url)
    main = soup.find("article") or soup.find("main") or soup.body or soup
    text = _clean(main.get_text("\n"))
    if len(text) < 200:
        raise HTTPException(
            status_code=422,
            detail="Cette page ne contient pas de texte exploitable. Les plateformes vidéo protègent "
                   "leurs transcriptions : exportez le sous-titre (.vtt/.srt) ou le fichier audio, puis importez-le.",
        )
    return {"title": title[:120], "text": text}
